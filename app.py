# app.py
# AI Property Registry & Loan Assistant (Lightweight, Streamlit Free-tier friendly)
# - Template-based "GenAI" registry generation (Jinja2)
# - In-app ML loan approval & fraud scoring (scikit-learn)
# - Webcam capture via st.camera_input (embedded into doc)
# - No heavy LLMs or CV libs to ensure Streamlit Free compatibility

import warnings
warnings.filterwarnings("ignore")

import streamlit as st
import pandas as pd
import numpy as np
from jinja2 import Template
from io import BytesIO
from datetime import datetime
import base64

# small ML libraries
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler

# optional docx export
from docx import Document

# ---- Page styling (simple, safe) ----
st.set_page_config(page_title="AI Property Registry & Loan Assistant", layout="centered")
st.markdown(
    """
    <style>
    .hero {
        background: linear-gradient(135deg,#4b6cb7,#182848);
        padding: 28px;
        border-radius: 12px;
        color: white;
        text-align:center;
        box-shadow: 0 8px 30px rgba(29,53,87,0.25);
    }
    .muted { color: #e6eef8; font-size:14px; }
    .card { background: white; padding:16px; border-radius:12px; box-shadow: 0 6px 18px rgba(23,43,77,0.06); }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown('<div class="hero"><h1>AI Property Registry & Loan Assistant</h1>'
            '<div class="muted">Generate property registry docs, check loan approval & flag fraud — beautiful, one-page demo.</div></div>',
            unsafe_allow_html=True)

st.markdown("")  # spacing

# ---- Sidebar: small help and how to use ----
st.sidebar.header("How to use (quick)")
st.sidebar.markdown(
    "1. Fill property + buyer details.\n"
    "2. Capture or upload buyer photo (optional).\n"
    "3. Click **Run Assessment** to see loan approval probability & fraud score.\n"
    "4. Click **Generate Registry Document** to download a registry draft with embedded photo."
)

# ---- Step 1: Forms for property + buyer + loan fields ----
st.subheader("1. Property & Buyer Details")

with st.form("details"):
    col1, col2 = st.columns(2)
    with col1:
        property_type = st.selectbox("Property Type", ["Flat / Apartment", "Land / Plot", "House / Villa", "Commercial"])
        property_address = st.text_input("Property Address", "123 Future Lane, Berlin")
        property_id = st.text_input("Property ID / Registry No.", value=f"PROP-{np.random.randint(10000,99999)}")
        property_area = st.number_input("Area (sq.m)", min_value=1, max_value=100000, value=85)
    with col2:
        buyer_name = st.text_input("Buyer Full Name", "Amit Sharma")
        buyer_dob = st.date_input("Buyer Date of Birth")
        buyer_nationality = st.text_input("Buyer Nationality", "Indian")
        buyer_email = st.text_input("Buyer Email", "amit@example.com")
    st.markdown("---")
    st.subheader("Loan / Financial Info (for approval check)")
    col3, col4 = st.columns(2)
    with col3:
        annual_income = st.number_input("Annual Income (USD)", min_value=0, value=45000)
        existing_debt = st.number_input("Existing Debt (USD)", min_value=0, value=5000)
        monthly_obligation = st.number_input("Monthly Debt Payments (USD)", min_value=0, value=300)
    with col4:
        requested_loan = st.number_input("Requested Loan Amount (USD)", min_value=100, value=120000)
        loan_term_years = st.number_input("Loan Term (years)", min_value=1, value=20)
        credit_score = st.slider("Credit Score (simulated)", 300, 850, 650)
    submitted = st.form_submit_button("Save Details")

if submitted:
    st.success("Details saved (for this session).")

# ---- Step 2: Webcam / Photo capture ----
st.subheader("2. Capture Buyer Photo (optional)")
st.markdown("Use your webcam to capture a photo that will be embedded into the registry document.")
img_file = st.camera_input("Take a real-time photo (or upload one)")

if img_file is not None:
    st.image(img_file, caption="Captured photo", width=220)
    # convert to base64 for embedding in generated doc
    img_bytes = img_file.getvalue()
    img_b64 = base64.b64encode(img_bytes).decode()

else:
    img_b64 = None

st.markdown("---")

# ---- Step 3: Lightweight ML model for loan approval and fraud detection ----

# Build a tiny synthetic dataset and train a logistic regression (done once per session)
@st.cache_resource
def train_demo_model():
    # synthetic dataset (simple features)
    rng = np.random.RandomState(42)
    n = 600
    income = rng.normal(50000, 20000, n).clip(5000, 300000)
    debt = rng.normal(8000, 6000, n).clip(0, 200000)
    monthly = (debt / 12) * rng.uniform(0.02, 0.12, n)
    credit = rng.normal(650, 70, n).clip(300, 850)
    loan_req = rng.normal(120000, 50000, n).clip(1000, 2000000)

    # target: approved if debt-to-income ratio small, credit high, and loan smaller than threshold
    dti = (debt + monthly * 12) / (income + 1)
    approved = ((dti < 0.4) & (credit > 600) & (loan_req < (income * 5))).astype(int)

    X = np.vstack([income, debt, monthly, credit, loan_req]).T
    scaler = StandardScaler().fit(X)
    Xs = scaler.transform(X)
    model = LogisticRegression(max_iter=300)
    model.fit(Xs, approved)
    return model, scaler

model, scaler = train_demo_model()

# risk scoring function (combines simple heuristics + model probability)
def assess_loan_and_fraud(annual_income, existing_debt, monthly_obligation, requested_loan, credit_score, buyer_age=None):
    features = np.array([[annual_income, existing_debt, monthly_obligation, credit_score, requested_loan]])
    Xs = scaler.transform(features)
    prob = model.predict_proba(Xs)[0,1]  # probability of approval
    # rule-based fraud heuristics (simple for demo)
    fraud_score = 0.0
    if requested_loan > annual_income * 7:
        fraud_score += 0.4
    if credit_score < 500:
        fraud_score += 0.3
    # age check (if provided): suspicious if age < 18
    if buyer_age is not None and buyer_age < 18:
        fraud_score += 0.3
    # normalize
    fraud_score = min(1.0, fraud_score)
    decision = "APPROVE" if prob > 0.5 and fraud_score < 0.5 else "REJECT"
    return prob, fraud_score, decision

# calculate buyer age
try:
    buyer_age = (datetime.now().date() - buyer_dob).days // 365
except Exception:
    buyer_age = None

# Button to run assessment
st.subheader("3. Run Assessment")
if st.button("Run Assessment"):
    prob, fraud_score, decision = assess_loan_and_fraud(
        annual_income, existing_debt, monthly_obligation, requested_loan, credit_score, buyer_age
    )
    st.metric("Loan Approval Probability", f"{prob*100:.1f}%")
    st.metric("Fraud Risk Score", f"{fraud_score*100:.1f}%")
    if decision == "APPROVE":
        st.success(f"Preliminary Decision: {decision}")
    else:
        st.error(f"Preliminary Decision: {decision}")

    # Explainability: what influenced the decision
    st.subheader("Why this decision (explainability)")
    reasons = []
    if prob < 0.5:
        reasons.append("- Low model approval probability based on income/debt/loan size.")
    else:
        reasons.append("- Model indicates reasonable debt-to-income and credit profile.")
    if fraud_score >= 0.5:
        reasons.append("- Fraud heuristics flagged risk (loan size vs income or low credit score).")
    else:
        reasons.append("- Fraud heuristics do not indicate major red flags.")
    for r in reasons:
        st.write(r)

st.markdown("---")

# ---- Step 4: Registry Document Generation (Template-based, Jinja2) ----

st.subheader("4. Generate Registry Document (downloadable)")

# A simple Jinja2 template for the registry document
REGISTRY_TEMPLATE = """
PROPERTY REGISTRY / TRANSFER DOCUMENT
------------------------------------
Registry ID: {{ property_id }}
Date: {{ date }}

Property Type: {{ property_type }}
Address: {{ property_address }}
Area (sq.m): {{ property_area }}

OWNER / BUYER:
Name: {{ buyer_name }}
Date of Birth: {{ buyer_dob }}
Nationality: {{ buyer_nationality }}
Email: {{ buyer_email }}

LOAN DETAILS:
Requested Loan: USD {{ requested_loan }}
Loan Term (years): {{ loan_term_years }}
Preliminary Decision: {{ decision }}
Loan Approval Probability: {{ prob_percent }}
Fraud Risk Score: {{ fraud_percent }}

AGENTIC NOTES:
- This document was generated using templated GenAI-style synthesis in a demo app.
- The buyer's photo (if provided) is embedded below.

SIGNATURE:
___________________________

(End of document)
"""

def render_registry_text(context: dict):
    tpl = Template(REGISTRY_TEMPLATE)
    return tpl.render(**context)

def make_docx(text: str, image_bytes=None):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    if image_bytes:
        # add a small image at the end
        doc.add_page_break()
        doc.add_paragraph("Embedded photo:")
        doc.add_picture(BytesIO(image_bytes), width=docx.shared.Inches(2))
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Prepare context
context = {
    "property_id": property_id,
    "date": datetime.now().strftime("%Y-%m-%d"),
    "property_type": property_type,
    "property_address": property_address,
    "property_area": property_area,
    "buyer_name": buyer_name,
    "buyer_dob": buyer_dob,
    "buyer_nationality": buyer_nationality,
    "buyer_email": buyer_email,
    "requested_loan": requested_loan,
    "loan_term_years": loan_term_years,
    "decision": "Not assessed" ,
    "prob_percent": "N/A",
    "fraud_percent": "N/A"
}

# Update context if assessment run
if 'prob' in locals():
    context["decision"] = decision
    context["prob_percent"] = f"{prob*100:.1f}%"
    context["fraud_percent"] = f"{fraud_score*100:.1f}%"

# Generate text
registry_text = render_registry_text(context)

st.text_area("Preview registry document (editable)", registry_text, height=300)

# Buttons to download
def download_text_file(text, filename="registry.txt"):
    b = text.encode("utf-8")
    b64 = base64.b64encode(b).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download TXT</a>'
    st.markdown(href, unsafe_allow_html=True)

def download_docx_file(text, img_bytes=None, filename="registry.docx"):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    if img_bytes:
        # try to add image safely
        try:
            doc.add_page_break()
            doc.add_paragraph("Embedded photo:")
            doc.add_picture(BytesIO(img_bytes), width=docx.shared.Inches(2))
        except Exception:
            pass
    bio = BytesIO()
    doc.save(bio)
    b = bio.getvalue()
    b64 = base64.b64encode(b).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download DOCX</a>'
    st.markdown(href, unsafe_allow_html=True)

colb1, colb2 = st.columns(2)
with colb1:
    if st.button("Generate & Download TXT"):
        download_text_file(registry_text, filename=f"{property_id}_registry.txt")
with colb2:
    if st.button("Generate & Download DOCX"):
        # DOCX generation is basic; we will include image if present
        bio = BytesIO()
        doc = Document()
        for line in registry_text.splitlines():
            doc.add_paragraph(line)
        if img_b64:
            try:
                img_bytes2 = base64.b64decode(img_b64)
                doc.add_page_break()
                doc.add_paragraph("Embedded photo:")
                doc.add_picture(BytesIO(img_bytes2), width=docx.shared.Inches(2))
            except Exception:
                pass
        doc.save(bio)
        b = bio.getvalue()
        b64 = base64.b64encode(b).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{property_id}_registry.docx">Download DOCX</a>'
        st.markdown(href, unsafe_allow_html=True)

st.markdown("---")

# ---- Final UX explanation about limitations (honest) ----
st.info(
    "Note: This is an MVP demo. For full production-grade GenAI document generation and secure facial recognition, "
    "integrate certified LLMs and secure biometric systems. This demo uses template generation and lightweight ML to stay fully free & deployable."
)
